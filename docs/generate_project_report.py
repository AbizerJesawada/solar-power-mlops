from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from PIL import Image, ImageDraw, ImageFont, ImageFilter


ROOT = Path(__file__).resolve().parents[1]
REPORTS_DIR = ROOT / "reports"
OUTPUT = REPORTS_DIR / "Solar_Power_MLOps_Project_Report.docx"
ASSETS_DIR = REPORTS_DIR / "report_assets"

REPO_URL = "https://github.com/AbizerJesawada/solar-power-mlops"
DEPLOYED_URL = "http://65.1.133.227:8501"

FONT_REGULAR_CANDIDATES = [
    Path("C:/Windows/Fonts/aptos.ttf"),
    Path("C:/Windows/Fonts/segoeui.ttf"),
    Path("C:/Windows/Fonts/calibri.ttf"),
    Path("C:/Windows/Fonts/arial.ttf"),
]

FONT_BOLD_CANDIDATES = [
    Path("C:/Windows/Fonts/aptos-bold.ttf"),
    Path("C:/Windows/Fonts/segoeuib.ttf"),
    Path("C:/Windows/Fonts/calibrib.ttf"),
    Path("C:/Windows/Fonts/arialbd.ttf"),
]


def ensure_assets_dir():
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)


def get_font(size: int, *, bold: bool = False):
    candidates = FONT_BOLD_CANDIDATES if bold else FONT_REGULAR_CANDIDATES
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def text_box(draw: ImageDraw.ImageDraw, text: str, font, max_width: int):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        test = f"{current} {word}".strip()
        if draw.textbbox((0, 0), test, font=font)[2] <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_multiline_center(draw, box, text, font, fill, spacing=8):
    lines = []
    for raw in text.split("\n"):
        lines.extend(text_box(draw, raw, font, box[2] - box[0] - 28))
    line_heights = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        line_heights.append(bbox[3] - bbox[1])
    total_height = sum(line_heights) + spacing * max(0, len(lines) - 1)
    y = box[1] + ((box[3] - box[1]) - total_height) / 2
    for idx, line in enumerate(lines):
        bbox = draw.textbbox((0, 0), line, font=font)
        line_width = bbox[2] - bbox[0]
        draw.text(
            (box[0] + ((box[2] - box[0]) - line_width) / 2, y),
            line,
            font=font,
            fill=fill,
        )
        y += line_heights[idx] + spacing


def draw_card(base, box, title, fill, edge, icon_fill=None):
    shadow = Image.new("RGBA", base.size, (0, 0, 0, 0))
    shadow_draw = ImageDraw.Draw(shadow)
    sx1, sy1, sx2, sy2 = box[0] + 10, box[1] + 12, box[2] + 10, box[3] + 12
    shadow_draw.rounded_rectangle((sx1, sy1, sx2, sy2), radius=28, fill=(51, 65, 85, 60))
    shadow = shadow.filter(ImageFilter.GaussianBlur(8))
    base.alpha_composite(shadow)

    draw = ImageDraw.Draw(base)
    draw.rounded_rectangle(box, radius=28, fill=fill, outline=edge, width=4)
    draw.line((box[0] + 24, box[1] + 22, box[2] - 24, box[1] + 22), fill=edge, width=5)
    if icon_fill:
        circle = (box[0] + 18, box[1] + 16, box[0] + 54, box[1] + 52)
        draw.ellipse(circle, fill=icon_fill)
    font = get_font(28, bold=True)
    draw_multiline_center(draw, box, title, font, "#0F172A", spacing=8)


def draw_dark_card(base, box, title, subtitle, fill, edge, title_color="#E5E7EB", subtitle_color="#C7D2FE"):
    draw = ImageDraw.Draw(base)
    draw.rounded_rectangle(box, radius=24, fill=fill, outline=edge, width=3)
    title_font = get_font(26, bold=True)
    subtitle_font = get_font(18)

    title_bbox = draw.textbbox((0, 0), title, font=title_font)
    title_y = box[1] + 18
    draw.text((box[0] + 40, title_y), title, font=title_font, fill=title_color)

    subtitle_lines = []
    for raw in subtitle.split("\n"):
        subtitle_lines.extend(text_box(draw, raw, subtitle_font, box[2] - box[0] - 70))
    y = title_y + (title_bbox[3] - title_bbox[1]) + 10
    for line in subtitle_lines:
        draw.text((box[0] + 40, y), line, font=subtitle_font, fill=subtitle_color)
        line_bbox = draw.textbbox((0, 0), line, font=subtitle_font)
        y += (line_bbox[3] - line_bbox[1]) + 6


def draw_line_arrow(draw, start, end, fill="#475569", width=6, arrow_size=18):
    draw.line([start, end], fill=fill, width=width)
    dx = end[0] - start[0]
    dy = end[1] - start[1]
    length = max((dx ** 2 + dy ** 2) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    tip = end
    left = (end[0] - ux * arrow_size - px * arrow_size * 0.55, end[1] - uy * arrow_size - py * arrow_size * 0.55)
    right = (end[0] - ux * arrow_size + px * arrow_size * 0.55, end[1] - uy * arrow_size + py * arrow_size * 0.55)
    draw.polygon([tip, left, right], fill=fill)


def draw_lane_band(draw, box, title, fill, edge):
    draw.rounded_rectangle(box, radius=32, fill=fill, outline=edge, width=3)
    label_box = (box[0] + 20, box[1] + 16, box[0] + 280, box[1] + 62)
    draw.rounded_rectangle(label_box, radius=18, fill=edge)
    label_font = get_font(22, bold=True)
    draw_multiline_center(draw, label_box, title, label_font, "white", spacing=4)


def create_canvas(width, height):
    return Image.new("RGBA", (width, height), "#F8FAFC")


def save_canvas(image, output):
    image.convert("RGB").save(output, quality=95)


def draw_lane(ax, xy, width, height, title, fill="#F8FAFC", edge="#D7DEE8"):
    lane = FancyBboxPatch(
        xy,
        width,
        height,
        boxstyle="round,pad=0.01,rounding_size=0.025",
        linewidth=1.0,
        edgecolor=edge,
        facecolor=fill,
        zorder=0,
    )
    ax.add_patch(lane)
    ax.text(
        xy[0] + 0.02,
        xy[1] + height - 0.035,
        title,
        ha="left",
        va="center",
        fontsize=11,
        color="#334155",
        weight="bold",
    )


def draw_box(ax, xy, width, height, text, facecolor="#E8F0FE", edgecolor="#3367D6", fontsize=10):
    shadow = FancyBboxPatch(
        (xy[0] + 0.008, xy[1] - 0.01),
        width,
        height,
        boxstyle="round,pad=0.02,rounding_size=0.02",
        linewidth=0,
        facecolor="#CBD5E1",
        alpha=0.35,
        zorder=1,
    )
    ax.add_patch(shadow)
    box = FancyBboxPatch(
        xy,
        width,
        height,
        boxstyle="round,pad=0.02,rounding_size=0.02",
        linewidth=1.8,
        edgecolor=edgecolor,
        facecolor=facecolor,
        zorder=2,
    )
    ax.add_patch(box)
    ax.plot(
        [xy[0] + 0.02, xy[0] + width - 0.02],
        [xy[1] + height - 0.02, xy[1] + height - 0.02],
        color=edgecolor,
        lw=2.2,
        solid_capstyle="round",
        zorder=3,
    )
    ax.text(
        xy[0] + width / 2,
        xy[1] + height / 2,
        text,
        ha="center",
        va="center",
        fontsize=fontsize,
        weight="bold",
        wrap=True,
        color="#0F172A",
        zorder=4,
    )


def draw_arrow(ax, start, end, text=None):
    ax.annotate(
        "",
        xy=end,
        xytext=start,
        arrowprops=dict(arrowstyle="-|>", lw=1.8, color="#475569", shrinkA=6, shrinkB=6),
    )
    if text:
        ax.text(
            (start[0] + end[0]) / 2,
            (start[1] + end[1]) / 2 + 0.02,
            text,
            fontsize=9,
            ha="center",
            va="center",
            color="#333333",
        )


def box_center(xy, width, height):
    return (xy[0] + width / 2, xy[1] + height / 2)


def generate_architecture_diagram() -> Path:
    ensure_assets_dir()
    output = ASSETS_DIR / "system_architecture_diagram.png"
    image = Image.new("RGBA", (1800, 1100), "#050505")
    draw = ImageDraw.Draw(image)

    title_font = get_font(52, bold=True)
    subtitle_font = get_font(22)
    draw.text((900, 52), "System architecture diagram", anchor="mm", font=title_font, fill="#F8FAFC")
    draw.text((900, 108), "End-to-end MLOps components and deployment path", anchor="mm", font=subtitle_font, fill="#D4D4D8")

    draw.rounded_rectangle((110, 170, 1690, 440), radius=34, fill="#0F5B4C", outline="#5EEAD4", width=2)
    draw.rounded_rectangle((110, 500, 1690, 770), radius=34, fill="#4C3F9B", outline="#A78BFA", width=2)
    draw.rounded_rectangle((110, 830, 1690, 1070), radius=34, fill="#565656", outline="#A3A3A3", width=2)

    lane_font = get_font(24, bold=True)
    draw.text((150, 205), "Data foundation", font=lane_font, fill="#0B7661")
    draw.text((150, 535), "Model and experiment", font=lane_font, fill="#5C50C4")
    draw.text((150, 865), "Delivery and cloud", font=lane_font, fill="#737373")

    cards = {
        "raw": (160, 245, 550, 385),
        "dvc": (640, 245, 1030, 385),
        "prep": (1120, 245, 1580, 385),
        "train": (160, 575, 550, 715),
        "mlflow": (640, 575, 1090, 715),
        "eval": (1170, 575, 1580, 715),
        "app": (160, 905, 600, 1015),
        "github": (700, 905, 1140, 1015),
        "ec2": (1210, 905, 1600, 1015),
    }

    draw_dark_card(image, cards["raw"], "Raw data", "Generation + weather", "#115E50", "#6EE7B7", "#B7F7DF", "#6EE7B7")
    draw_dark_card(image, cards["dvc"], "DVC + AWS S3", "Data versioning", "#2F5D0C", "#A3E635", "#D9F99D", "#A3E635")
    draw_dark_card(image, cards["prep"], "Ingestion + preprocessing", "Feature engineering", "#1E4F8A", "#93C5FD", "#DBEAFE", "#93C5FD")
    draw_dark_card(image, cards["train"], "Model training", "XGBoost regressor", "#4A3FA3", "#C4B5FD", "#DDD6FE", "#C4B5FD")
    draw_dark_card(image, cards["mlflow"], "MLflow tracking", "Params, metrics, artifacts", "#7C4706", "#FBBF24", "#FDE68A", "#FBBF24")
    draw_dark_card(image, cards["eval"], "Evaluation + monitoring", "Plots + drift report", "#8A3817", "#FDBA74", "#FED7AA", "#FDBA74")
    draw_dark_card(image, cards["app"], "Streamlit app", "Real-time prediction", "#1E4F8A", "#93C5FD", "#DBEAFE", "#93C5FD")
    draw_dark_card(image, cards["github"], "GitHub Actions", "CI/CD automation", "#57534E", "#D6D3D1", "#E7E5E4", "#D6D3D1")
    draw_dark_card(image, cards["ec2"], "AWS EC2", "Cloud hosting", "#115E50", "#5EEAD4", "#99F6E4", "#5EEAD4")

    def center_right(box):
        return (box[2], (box[1] + box[3]) // 2)

    def center_left(box):
        return (box[0], (box[1] + box[3]) // 2)

    def center_bottom(box):
        return ((box[0] + box[2]) // 2, box[3])

    def center_top(box):
        return ((box[0] + box[2]) // 2, box[1])

    draw_line_arrow(draw, center_right(cards["raw"]), center_left(cards["dvc"]), fill="#A3A3A3", width=5)
    draw_line_arrow(draw, center_right(cards["dvc"]), center_left(cards["prep"]), fill="#A3A3A3", width=5)
    draw_line_arrow(draw, center_bottom(cards["raw"]), center_top(cards["train"]), fill="#A3A3A3", width=5)
    draw_line_arrow(draw, center_bottom(cards["prep"]), center_top(cards["eval"]), fill="#A3A3A3", width=5)
    draw_line_arrow(draw, center_right(cards["train"]), center_left(cards["mlflow"]), fill="#A3A3A3", width=5)
    draw_line_arrow(draw, center_right(cards["mlflow"]), center_left(cards["eval"]), fill="#A3A3A3", width=5)
    draw_line_arrow(draw, center_bottom(cards["train"]), center_top(cards["app"]), fill="#A3A3A3", width=5)
    draw_line_arrow(draw, center_bottom(cards["eval"]), center_top(cards["ec2"]), fill="#A3A3A3", width=5)
    draw_line_arrow(draw, center_right(cards["app"]), center_left(cards["github"]), fill="#A3A3A3", width=5)
    draw_line_arrow(draw, center_right(cards["github"]), center_left(cards["ec2"]), fill="#A3A3A3", width=5)

    save_canvas(image, output)
    return output


def generate_pipeline_flow_diagram() -> Path:
    ensure_assets_dir()
    output = ASSETS_DIR / "pipeline_flow_diagram.png"
    image = Image.new("RGBA", (1500, 1900), "#050505")
    draw = ImageDraw.Draw(image)

    title_font = get_font(46, bold=True)
    subtitle_font = get_font(22)
    draw.text((750, 54), "End-to-end pipeline flow", anchor="mm", font=title_font, fill="#F8FAFC")
    draw.text((750, 104), "Operational sequence from versioned data to cloud-hosted inference", anchor="mm", font=subtitle_font, fill="#D4D4D8")

    draw.line((210, 245, 210, 1710), fill="#D4D4D8", width=6)
    draw.line((210, 1710, 210, 1800), fill="#D4D4D8", width=6)

    steps = [
        ("1", "Data versioning", "DVC + S3 remote stores reproducible raw data snapshots", "#065F46", "#34D399"),
        ("2", "Ingestion", "data_ingestion.py loads generation and weather datasets", "#1D4ED8", "#93C5FD"),
        ("3", "Preprocessing", "preprocessing.py merges data and creates HOUR, DAY, MONTH features", "#4338CA", "#C4B5FD"),
        ("4", "Training", "train.py fits the XGBoost forecasting model with MLflow logging", "#A16207", "#F59E0B"),
        ("5", "Evaluation", "evaluate.py computes RMSE, MAE, R², and result visualizations", "#9A3412", "#FB923C"),
        ("6", "Monitoring", "monitoring.py generates drift report and governance outputs", "#3F6212", "#84CC16"),
        ("7", "Prediction app", "app.py exposes real-time AC power prediction via Streamlit", "#1D4ED8", "#93C5FD"),
        ("8", "Deployment", "Application runs on AWS EC2 using Docker and CI/CD validation", "#525252", "#A3A3A3"),
    ]

    top = 250
    row_gap = 170
    card_box = (300, 0, 1320, 110)
    for idx, (num, title, body, accent, fill) in enumerate(steps):
        y = top + idx * row_gap
        draw.ellipse((160, y + 6, 260, y + 106), fill=accent, outline="#E5E7EB", width=2)
        num_font = get_font(34, bold=True)
        draw.text((210, y + 56), num, anchor="mm", font=num_font, fill="#E5E7EB")
        box = (card_box[0], y, card_box[2], y + 110)
        draw_dark_card(image, box, title, body, accent, fill, "#E5E7EB", fill)
        draw_line_arrow(draw, (260, y + 56), (300, y + 56), fill="#D4D4D8", width=4)
        if idx < len(steps) - 1:
            draw_line_arrow(draw, (210, y + 106), (210, y + row_gap + 6), fill="#D4D4D8", width=4)

    save_canvas(image, output)
    return output


def generate_methodology_diagram() -> Path:
    ensure_assets_dir()
    output = ASSETS_DIR / "methodology_diagram.png"
    image = create_canvas(1800, 1100)
    draw = ImageDraw.Draw(image)

    title_font = get_font(44, bold=True)
    subtitle_font = get_font(20)
    draw.text((900, 58), "Project Methodology Diagram", anchor="mm", font=title_font, fill="#1F2937")
    draw.text((900, 108), "How the project moves from data preparation to monitored deployment", anchor="mm", font=subtitle_font, fill="#6B7280")

    draw.rounded_rectangle((80, 180, 1720, 520), radius=34, fill="#FCFAFF", outline="#C4B5FD", width=2)
    draw.rounded_rectangle((80, 620, 1720, 960), radius=34, fill="#FBFEFB", outline="#CDECCF", width=2)

    tag_font = get_font(18, bold=True)
    draw.rounded_rectangle((100, 200, 360, 245), radius=18, fill="#E9D5FF")
    draw.text((230, 223), "Core Development\nPhases", anchor="mm", font=tag_font, fill="white")
    draw.rounded_rectangle((100, 640, 360, 685), radius=18, fill="#D9F99D")
    draw.text((230, 663), "Operational\nOutcomes", anchor="mm", font=tag_font, fill="white")

    top_boxes = [
        ((140, 295, 520, 435), "Data Handling\nLoad, merge, clean,\ncreate time features", "#F8FBFF", "#2563EB", "#93C5FD"),
        ((700, 295, 1080, 435), "Model Development\nTrain/test split,\nXGBoost training", "#FCF8FF", "#9333EA", "#D8B4FE"),
        ((1260, 295, 1640, 435), "Experiment Strategy\nDVC + MLflow +\nparams.yaml", "#F7FFF9", "#16A34A", "#86EFAC"),
    ]
    bottom_boxes = [
        ((140, 735, 520, 875), "Evaluation\nRMSE, MAE, R2,\nplots", "#FFF8FB", "#DB2777", "#F9A8D4"),
        ((700, 735, 1080, 875), "Monitoring\nDrift report,\npipeline logs", "#F5FDFF", "#0891B2", "#A5F3FC"),
        ((1260, 735, 1640, 875), "Deployment\nStreamlit, Docker,\nAWS EC2", "#FFF9F5", "#EA580C", "#FDBA74"),
    ]

    for box, title, fill, edge, icon in top_boxes + bottom_boxes:
        draw_card(image, box, title, fill, edge, icon)

    draw_line_arrow(draw, (520, 365), (700, 365), fill="#475569", width=4)
    draw_line_arrow(draw, (1080, 365), (1260, 365), fill="#475569", width=4)
    draw_line_arrow(draw, (330, 435), (330, 735), fill="#475569", width=4)
    draw_line_arrow(draw, (890, 435), (890, 735), fill="#475569", width=4)
    draw_line_arrow(draw, (1450, 435), (1450, 735), fill="#475569", width=4)

    save_canvas(image, output)
    return output


def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_base_style(document: Document):
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.2

    for name, size in [("Title", 20), ("Heading 1", 16), ("Heading 2", 14)]:
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        if style.font:
            style.font.color.rgb = RGBColor(31, 41, 55)


def add_paragraph(document: Document, text: str, *, bold: bool = False, italic: bool = False):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def add_bullet(document: Document, text: str):
    document.add_paragraph(text, style="List Bullet")


def add_table(document: Document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, value in enumerate(headers):
        hdr[i].text = str(value)
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = "Times New Roman"
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "1F4E79")
            hdr[i]._tc.get_or_add_tcPr().append(shading)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    return table


def add_image(document: Document, path: Path, caption: str, width: float = 5.8):
    if not path.exists():
        return
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = document.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    if cap.runs:
        cap.runs[0].italic = True
        cap.runs[0].font.color.rgb = RGBColor(71, 85, 105)


def add_cover_block(document: Document):
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(70)
    title.paragraph_format.space_after = Pt(20)
    run = title.add_run("Solar Power Generation Forecasting")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(20, 58, 95)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run("An End-to-End MLOps Pipeline on AWS")
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(71, 85, 105)

    for line in [
        "Student Names: _________________________________",
        "Roll Numbers: _________________________________",
        "Course Name: Essentials of MLOps",
        "Faculty Name: _________________________________",
        "Department of Artificial Intelligence and Machine Learning",
        "Academic Year: 2025-2026",
    ]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_section_divider(document: Document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 55)
    run.font.color.rgb = RGBColor(148, 163, 184)


def add_section_heading(document: Document, title: str, level: int = 1):
    heading = document.add_heading(title, level=level)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    if heading.runs:
        heading.runs[0].font.color.rgb = RGBColor(31, 41, 55)
    return heading


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def architecture_points():
    return [
        "Raw generation and weather datasets are tracked by DVC and stored remotely on AWS S3.",
        "The ingestion layer loads both datasets and validates their shapes before processing.",
        "The preprocessing layer merges the datasets on DATE_TIME, creates time-based features, and writes the processed file used for modeling.",
        "The training layer builds an XGBoost regression model and logs parameters plus model artifacts to MLflow.",
        "The evaluation layer computes RMSE, MAE, and R2 Score, and generates evaluation visualizations.",
        "The monitoring layer measures feature drift and generates drift reports for governance.",
        "The deployment layer exposes the trained model through Streamlit and Docker, and the application is hosted on AWS EC2.",
        "The CI/CD layer validates the end-to-end workflow through GitHub Actions by pulling DVC data, running the pipeline, and checking the app.",
    ]


def add_contents_block(document: Document):
    add_section_heading(document, "Report Outline")
    for item in [
        "1. Abstract",
        "2. Introduction",
        "3. Literature Review",
        "4. System Architecture",
        "5. Pipeline Flow",
        "6. Tools and Technologies Used",
        "7. Methodology",
        "8. Experiments and Results",
        "9. Comparative Study",
        "10. Deployment",
        "11. Monitoring and Maintenance",
        "12. Conclusion and Future Work",
        "13. References",
        "14. GitHub Repository",
        "15. Appendix",
    ]:
        add_bullet(document, item)


def write_report():
    metrics = load_json(REPORTS_DIR / "metrics.json")
    drift = load_json(REPORTS_DIR / "drift_report.json")
    architecture_diagram = generate_architecture_diagram()
    pipeline_diagram = generate_pipeline_flow_diagram()
    methodology_diagram = generate_methodology_diagram()

    document = Document()
    set_base_style(document)

    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_page_number(section)

    add_cover_block(document)

    document.add_page_break()
    add_contents_block(document)
    add_section_divider(document)

    add_section_heading(document, "Abstract")
    add_paragraph(
        document,
        "This project implements an end-to-end MLOps pipeline for forecasting solar power generation using weather and time-based features. "
        "The pipeline covers data versioning with DVC and AWS S3, experiment tracking with MLflow, modular training and evaluation, CI/CD with GitHub Actions, monitoring and governance, Docker containerization, and deployment on AWS EC2. "
        f"The final XGBoost regression model predicts AC_POWER and achieved RMSE {metrics['RMSE']}, MAE {metrics['MAE']}, and R2 Score {metrics['R2_Score']}. "
        "A Streamlit application was built for real-time inference, drift monitoring was implemented for governance, and the deployed Dockerized application was verified on an EC2 Ubuntu instance."
    )
    add_section_heading(document, "Executive Summary", 2)
    add_table(
        document,
        ["Area", "Summary"],
        [
            ["Business Use Case", "Forecast solar AC power generation from weather and time features."],
            ["Model", "XGBoost Regressor trained on six forecasting-oriented input features."],
            ["MLOps Stack", "DVC, AWS S3, MLflow, GitHub Actions, Streamlit, Docker, AWS EC2."],
            ["Best Result", f"RMSE {metrics['RMSE']}, MAE {metrics['MAE']}, R2 {metrics['R2_Score']}."],
            ["Deployment Mode", "Real-time web application deployed on AWS EC2."],
            ["Monitoring", "Drift report, pipeline logs, and MLflow experiment history."],
        ],
    )
    add_section_divider(document)

    add_section_heading(document, "Introduction")
    add_section_heading(document, "Problem Statement", 2)
    add_paragraph(
        document,
        "Solar power generation is highly dependent on environmental conditions such as solar irradiation, ambient temperature, module temperature, and time of day. "
        "In renewable energy systems, forecasting the generated power is important for grid planning, energy management, and operational monitoring. "
        "This project addresses the problem of forecasting AC power generation from a solar plant using a machine learning model and an end-to-end MLOps pipeline."
    )
    add_section_heading(document, "Objectives", 2)
    for item in [
        "Develop a regression model to forecast solar AC power output.",
        "Build a modular and reproducible MLOps pipeline.",
        "Version raw data using DVC and store it in an AWS S3 remote.",
        "Track experiments, models, metrics, plots, and drift outputs using MLflow.",
        "Deploy the model as a real-time prediction application using Streamlit, Docker, and AWS EC2.",
        "Add CI/CD, monitoring, logging, and governance controls to improve reliability and maintainability.",
    ]:
        add_bullet(document, item)
    add_section_heading(document, "Scope of the Project", 2)
    add_paragraph(
        document,
        "The scope of this project includes raw data versioning, data preprocessing, model training, evaluation, monitoring, deployment, and governance documentation. "
        "The project focuses on a single plant dataset and a real-time prediction interface. It does not yet include HTTPS, automated production deployment, or advanced retraining orchestration."
    )
    add_section_heading(document, "Project Significance", 2)
    add_paragraph(
        document,
        "This project is significant because it demonstrates that a machine learning model can be treated as a full operational system rather than an isolated notebook artifact. "
        "It combines forecasting accuracy with reproducibility, deployment readiness, cloud execution, and governance practices expected in modern MLOps workflows."
    )
    add_section_divider(document)

    add_section_heading(document, "Literature Review")
    add_paragraph(
        document,
        "Forecasting renewable energy output has become an important research area because energy generation from solar plants is inherently variable. "
        "Most practical solar forecasting systems rely on historical plant data, weather measurements, and temporal features. Supervised regression algorithms are widely applied because they can learn the relationship between environmental conditions and generated power."
    )
    add_paragraph(
        document,
        "Among tabular machine learning methods, gradient boosting models such as XGBoost are frequently used because they can capture nonlinear dependencies, interactions between features, and irregular patterns in operational data. "
        "At the same time, modern MLOps practices emphasize that a model alone is not sufficient. A full production-style workflow requires data versioning, experiment tracking, deployment readiness, monitoring, logging, and governance controls."
    )
    add_paragraph(
        document,
        "This project adopts that broader MLOps perspective by combining a forecasting model with DVC, MLflow, GitHub Actions, Docker, Streamlit, AWS S3, and AWS EC2. "
        "The result is not only a trained model but a complete pipeline that is reproducible, auditable, deployable, and easier to maintain."
    )
    add_section_divider(document)

    add_section_heading(document, "System Architecture")
    add_section_heading(document, "Overall Architecture", 2)
    add_image(document, architecture_diagram, "Figure A. System Architecture Diagram", width=6.2)
    for point in architecture_points():
        add_bullet(document, point)
    add_paragraph(
        document,
        "The architecture separates responsibilities across data, modeling, experimentation, deployment, and governance. "
        "This separation makes the project easier to debug, extend, document, and evaluate against the Essentials of MLOps rubric."
    )
    add_section_divider(document)

    add_section_heading(document, "Pipeline Flow")
    add_image(document, pipeline_diagram, "Figure B. End-to-End Pipeline Flow Diagram", width=5.5)
    add_paragraph(
        document,
        "The end-to-end pipeline begins with DVC-managed raw CSV files. The generation dataset and weather dataset are ingested, cleaned, merged, and transformed into a processed modeling dataset. "
        "The processed dataset is then used to train an XGBoost model, evaluate prediction quality, generate plots, and run drift checks. "
        "After training, the model is exposed through a Streamlit application for real-time prediction. The same project is validated in CI and containerized for deployment."
    )
    for step in [
        "Data collection and versioning",
        "Data ingestion",
        "Preprocessing and feature engineering",
        "Model training",
        "Evaluation and report generation",
        "Drift monitoring",
        "Real-time prediction interface",
        "Containerization and cloud deployment",
        "CI/CD validation and governance",
    ]:
        add_bullet(document, step)
    add_section_divider(document)

    add_section_heading(document, "Tools and Technologies Used")
    add_table(
        document,
        ["Category", "Tools / Platforms", "Purpose"],
        [
            ["Programming", "Python", "Core implementation language"],
            ["Data Processing", "Pandas, NumPy", "Data loading, cleaning, merge, and feature handling"],
            ["Modeling", "XGBoost, Scikit-learn", "Regression and evaluation metrics"],
            ["Experiment Tracking", "MLflow", "Runs, parameters, metrics, model artifacts, drift artifacts"],
            ["Data Versioning", "DVC, AWS S3", "Raw data versioning and remote storage"],
            ["Visualization", "Matplotlib, Seaborn", "Exploratory and evaluation visualizations"],
            ["Deployment", "Streamlit, Docker, AWS EC2", "Real-time app, containerization, cloud hosting"],
            ["Automation", "GitHub Actions", "CI/CD validation workflow"],
            ["Governance and Logging", "Python logging, governance.md, Git", "Execution audit trail and process controls"],
        ],
    )
    add_section_heading(document, "Technology Selection Justification", 2)
    add_paragraph(
        document,
        "The technology stack was chosen to satisfy both technical practicality and rubric coverage. "
        "DVC and AWS S3 support data lineage, MLflow supports experiment accountability, GitHub Actions validates the pipeline in CI, Streamlit provides a fast real-time interface, Docker standardizes execution, and AWS EC2 offers an understandable and cost-conscious deployment path."
    )
    add_section_divider(document)

    add_section_heading(document, "Methodology")
    add_image(document, methodology_diagram, "Figure C. Methodology Diagram", width=6.2)
    add_section_heading(document, "Data Handling", 2)
    add_paragraph(
        document,
        "The project uses two raw datasets for Plant 1: generation data and weather sensor data. The generation data contains timestamped power and yield values, while the weather data contains irradiation and temperature measurements. "
        "During preprocessing, both datasets are aligned using DATE_TIME and merged into a single processed file."
    )
    add_paragraph(
        document,
        "The preprocessing stage converts timestamps, engineers HOUR, DAY, and MONTH features, removes null values, and writes the final dataset to data/processed/final_data.csv. "
        "Although the raw merged data contains other generation-related columns, the final forecasting model uses only weather and time features in order to avoid target leakage."
    )
    add_table(
        document,
        ["Dataset", "Rows", "Columns", "Role"],
        [
            ["Plant_1_Generation_Data.csv", "68,778", "7", "Historical generation measurements"],
            ["Plant_1_Weather_Sensor_Data.csv", "3,182", "6", "Environmental context and irradiation"],
            ["final_data.csv", "68,774", "10", "Merged and engineered modeling dataset"],
        ],
    )
    add_section_heading(document, "Model Development", 2)
    add_paragraph(
        document,
        "The target variable is AC_POWER. The model input features are AMBIENT_TEMPERATURE, MODULE_TEMPERATURE, IRRADIATION, HOUR, DAY, and MONTH. "
        "An XGBoost Regressor was selected because it performs strongly on structured tabular data and handles nonlinear relationships well."
    )
    add_paragraph(
        document,
        "The train-test split uses 80 percent training data and 20 percent testing data with random_state 42. Hyperparameters were configured through params.yaml so that the experiment remains reproducible and centrally configurable."
    )
    add_section_heading(document, "Versioning and Experimentation Strategy", 2)
    add_paragraph(
        document,
        "The project uses DVC to track raw data files and store them in an S3 remote. This ensures that raw data changes are reproducible and can be restored using dvc pull. "
        "MLflow is used to record experiment metadata, selected features, hyperparameters, train-test counts, model artifact, evaluation metrics, plots, and monitoring outputs."
    )
    add_section_heading(document, "Reproducibility Strategy", 2)
    add_paragraph(
        document,
        "Reproducibility is enforced through version-controlled code in Git, raw-data lineage in DVC, experiment history in MLflow, centralized configuration in params.yaml, and containerization in Docker. "
        "This combination reduces ambiguity when rerunning experiments and makes the project easier to defend during evaluation."
    )
    add_section_divider(document)

    add_section_heading(document, "Experiments and Results")
    add_section_heading(document, "Experimental Setup", 2)
    add_table(
        document,
        ["Item", "Value"],
        [
            ["Target Variable", "AC_POWER"],
            ["Features", "AMBIENT_TEMPERATURE, MODULE_TEMPERATURE, IRRADIATION, HOUR, DAY, MONTH"],
            ["Model", "XGBoost Regressor"],
            ["Test Size", "0.2"],
            ["Random State", "42"],
            ["n_estimators", "100"],
            ["max_depth", "6"],
            ["learning_rate", "0.1"],
        ],
    )
    add_section_heading(document, "Performance Metrics", 2)
    add_table(
        document,
        ["Metric", "Value", "Meaning"],
        [
            ["RMSE", metrics["RMSE"], "Root mean squared error"],
            ["MAE", metrics["MAE"], "Mean absolute error"],
            ["R2 Score", metrics["R2_Score"], "Explained variance score"],
        ],
    )
    add_paragraph(
        document,
        "The final model achieved a high R2 score while still relying only on weather and time features. This makes the forecasting setup more realistic than a model that uses generation-derived variables such as DC_POWER. "
        "Residual analysis and actual-versus-predicted plots show that the model captures the target trend effectively."
    )
    add_section_heading(document, "Interpretation of Results", 2)
    add_paragraph(
        document,
        "These results indicate that the selected feature set preserves strong predictive signal even after removing leakage-prone variables. "
        "The model therefore remains both accurate and defensible as a true forecasting solution rather than an overfitted shortcut."
    )

    add_section_heading(document, "Result Visualizations", 2)
    for path, caption in [
        (REPORTS_DIR / "actual_vs_predicted.png", "Figure 1. Actual vs Predicted AC Power"),
        (REPORTS_DIR / "residual_distribution.png", "Figure 2. Residual Distribution"),
        (REPORTS_DIR / "residuals_vs_predicted.png", "Figure 3. Residuals vs Predicted AC Power"),
        (REPORTS_DIR / "evaluation_metrics.png", "Figure 4. Evaluation Metrics"),
        (REPORTS_DIR / "correlation_heatmap.png", "Figure 5. Feature Correlation Heatmap"),
        (REPORTS_DIR / "hourly_power.png", "Figure 6. Average Solar Power by Hour"),
    ]:
        add_image(document, path, caption)
    add_section_divider(document)

    add_section_heading(document, "Comparative Study")
    add_paragraph(
        document,
        "An earlier intermediate version of the project used generation-related variables such as DC_POWER as inputs while predicting AC_POWER. Although that approach gave near-perfect results, it was not a strong forecasting design because those variables are too closely related to the target. "
        "The final model uses only weather and time-based features, which makes the prediction task more realistic, defensible, and suitable for demonstration in an MLOps context."
    )
    add_table(
        document,
        ["Approach", "Advantages", "Limitations"],
        [
            ["Leakage-prone feature set", "Artificially high scores", "Not suitable for realistic forecasting"],
            ["Final weather/time feature set", "Better reflects real deployment scenario", "Slightly harder prediction problem"],
            ["Implemented MLOps version", "Tracks data, experiments, deployment, and monitoring", "Future work can add automated retraining"],
        ],
    )
    add_section_divider(document)

    add_section_heading(document, "Deployment")
    add_section_heading(document, "Deployment Strategy", 2)
    add_paragraph(
        document,
        "A real-time deployment approach was selected for this project. The user provides current weather and time values through the Streamlit interface, and the model immediately returns the predicted AC power value. "
        "This approach is appropriate because the use case is interactive forecasting rather than offline batch scoring."
    )
    add_section_heading(document, "Deployment Implementation", 2)
    add_paragraph(
        document,
        "The Streamlit application was containerized using Docker and deployed on an AWS EC2 Ubuntu instance. "
        "Docker provides environment consistency, while EC2 offers flexible cloud hosting for the containerized app. The EC2 security group was configured to allow inbound traffic on port 8501."
    )
    add_table(
        document,
        ["Deployment Item", "Value"],
        [
            ["Platform", "AWS EC2 (Ubuntu 22.04)"],
            ["Instance Type", "t3.micro"],
            ["Container Runtime", "Docker"],
            ["Application Port", "8501"],
            ["Public URL", DEPLOYED_URL],
        ],
    )
    add_section_heading(document, "Deployment Rationale", 2)
    add_paragraph(
        document,
        "AWS EC2 was selected because it provides a simpler and more cost-aware deployment route than a fully managed serving platform for a mini project. "
        "This decision still satisfies the cloud deployment requirement while remaining easy to explain, demonstrate, and control."
    )
    add_section_divider(document)

    add_section_heading(document, "Monitoring and Maintenance")
    add_paragraph(
        document,
        "The monitoring stage compares reference and current feature distributions and writes a JSON drift report. "
        f"The configured drift threshold is {drift['threshold_percent']} percent. The latest run reported drift_detected = {drift['drift_detected']}. "
        "These monitoring outputs are also logged to MLflow, making them visible in the experiment history."
    )
    add_table(
        document,
        ["Feature", "Reference Mean", "Current Mean", "Drift %", "Drift Detected"],
        [
            [
                feature,
                values["reference_mean"],
                values["current_mean"],
                values["drift_percent"],
                values["drift_detected"],
            ]
            for feature, values in drift["features"].items()
        ],
    )
    add_paragraph(
        document,
        "Maintenance strategy for the project includes checking logs in logs/pipeline.log, reviewing MLflow runs, observing drift behavior, and retraining the model if drift or performance degradation becomes significant."
    )
    add_section_heading(document, "Governance Controls", 2)
    add_paragraph(
        document,
        "Governance is supported through several linked controls: DVC for dataset traceability, MLflow for experiment auditability, GitHub Actions for automated validation, Python logging for execution history, and drift reporting for operational oversight. "
        "These controls improve transparency, accountability, and long-term maintainability."
    )
    add_section_divider(document)

    add_section_heading(document, "Conclusion and Future Work")
    add_paragraph(
        document,
        "This project successfully demonstrates a complete MLOps workflow for solar power forecasting. It covers data versioning, modular pipeline design, experiment tracking, CI/CD, real-time deployment, monitoring, logging, governance, and cloud hosting. "
        "The final result is not just a machine learning model but a project that is reproducible, testable, deployable, and aligned with the Essentials of MLOps rubric."
    )
    add_paragraph(
        document,
        "Future work can include adding HTTPS and custom domain support for the deployed application, setting up automated deployment from GitHub Actions to AWS, introducing more advanced statistical drift detection, and implementing a retraining pipeline triggered by drift or performance decay."
    )
    add_section_divider(document)

    add_section_heading(document, "References")
    references = [
        '[1] T. Chen and C. Guestrin, "XGBoost: A Scalable Tree Boosting System," Proc. 22nd ACM SIGKDD Int. Conf. Knowledge Discovery and Data Mining, 2016, pp. 785-794.',
        '[2] MLflow, "MLflow Documentation." [Online]. Available: https://mlflow.org/docs/latest/index.html',
        '[3] DVC, "Data Version Control Documentation." [Online]. Available: https://dvc.org/doc',
        '[4] Docker, "Docker Documentation." [Online]. Available: https://docs.docker.com/',
        '[5] Streamlit, "Streamlit Documentation." [Online]. Available: https://docs.streamlit.io/',
        '[6] GitHub, "GitHub Actions Documentation." [Online]. Available: https://docs.github.com/actions',
        '[7] Amazon Web Services, "Amazon EC2 Documentation." [Online]. Available: https://docs.aws.amazon.com/ec2/',
    ]
    for ref in references:
        add_paragraph(document, ref)

    add_section_heading(document, "GitHub Repository")
    add_paragraph(document, REPO_URL)
    add_paragraph(
        document,
        "The repository contains the full project implementation, including source code, pipeline modules, configuration files, DVC pointers, MLflow tracking outputs, governance documentation, CI workflow, Docker files, and deployment instructions."
    )
    add_section_heading(document, "Rubric Coverage Summary", 2)
    add_table(
        document,
        ["Rubric Area", "Implementation Evidence"],
        [
            ["Problem Definition & ML Use Case", "Solar AC power forecasting with a clear regression objective"],
            ["Data Versioning & Experiment Tracking", "DVC + AWS S3 + MLflow"],
            ["Modular Pipeline Design", "Separate ingestion, preprocessing, training, evaluation, monitoring modules"],
            ["CI/CD", "GitHub Actions workflow"],
            ["Model Deployment Strategy", "Real-time Streamlit application"],
            ["Cloud Deployment & Infrastructure", "AWS EC2, AWS S3, Docker"],
            ["Monitoring, Logging & Governance", "Drift report, pipeline logs, governance document"],
            ["GitHub Repository & Reproducibility", "Structured repo, README, Docker, params, DVC, CI"],
        ],
    )
    add_section_divider(document)

    add_section_heading(document, "Appendix: Suggested Screenshots for Final Submission")
    for item in [
        "GitHub repository home page",
        "GitHub Actions successful workflow run",
        "MLflow experiment UI showing the latest unified run",
        "Streamlit prediction app screenshot",
        "AWS S3 bucket used as the DVC remote",
        "AWS EC2 instance details",
        "EC2 security group inbound rules for port 8501",
        "Docker container running on EC2 (`docker ps` output)",
        "Drift report JSON file",
        "Pipeline log file (`logs/pipeline.log`)",
    ]:
        add_bullet(document, item)

    appendix = document.add_section(WD_SECTION.NEW_PAGE)
    add_page_number(appendix)

    try:
        document.save(OUTPUT)
        final_output = OUTPUT
    except PermissionError:
        final_output = REPORTS_DIR / "Solar_Power_MLOps_Project_Report_Updated.docx"
        document.save(final_output)
    print(f"Generated {final_output}")


if __name__ == "__main__":
    write_report()
